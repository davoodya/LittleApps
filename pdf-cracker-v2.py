"""
OCR + Searchable PDF + DOCX generator (uses tesseract CLI via subprocess)

Requirements:
 - Tesseract OCR installed and available (set TESSERACT_CMD if not in PATH)
 - poppler (for pdf2image)
 - pip install pdf2image python-docx pillow opencv-python pypdf
"""

import os
import subprocess
from pathlib import Path
from pdf2image import convert_from_path
from PIL import Image
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
import shutil
import glob

# ---------- PDF File Input ----------
from pathlib import Path
import os

while True:
    path_input = input("Please Enter PDF Path: ").strip()
    PDF_INPUT = Path(path_input)

    if PDF_INPUT.exists():
        break
    else:
        print("[x] PDF Not Found, Enter Correct Path Please.")

# ---------- CONFIG ----------
OUT_PREFIX = "output"
DPI = 300
LANG = "eng"      # change to "fas" or "eng+fas" if you have Persian trained data
USE_DESKEW = True
POPPLER_PATH = r"H:\Repo\LittleApps\Materials\poppler\Library\bin"
# If tesseract executable is not in PATH, set full path, otherwise leave as "tesseract"
TESSERACT_CMD = None  # e.g. r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# Optional tesseract config flags:
TESSERACT_OEM = "3"
TESSERACT_PSM = "3"
# ----------------------------

tesseract_exe = TESSERACT_CMD or "tesseract"

# check tesseract existence
def check_tesseract():
    try:
        subprocess.run([tesseract_exe, "--version"], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except Exception as e:
        raise RuntimeError(f"Cannot run tesseract executable '{tesseract_exe}'. Ensure tesseract is installed and path is correct. Error: {e}")

check_tesseract()

work_dir = Path("ocr_temp")
work_dir.mkdir(exist_ok=True)

# helper: simple deskew and denoise using OpenCV
def preprocess_image(pil_image):
    # convert to OpenCV format
    img = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    # grayscale
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    # denoise
    denoised = cv2.fastNlMeansDenoising(gray, None, h=10, templateWindowSize=7, searchWindowSize=21)
    # threshold (adaptive)
    th = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                               cv2.THRESH_BINARY, 41, 11)
    if USE_DESKEW:
        coords = cv2.findNonZero(255 - th)
        if coords is not None:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            (h, w) = th.shape
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            rotated = cv2.warpAffine(th, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            th = rotated
    return Image.fromarray(th)

print("Converting PDF to images...")
pages = convert_from_path(PDF_INPUT, dpi=DPI, poppler_path=POPPLER_PATH)

doc = Document()
all_texts = []

# store per-page pdf paths for merging
page_pdf_paths = []

for i, page in enumerate(pages, start=1):
    print(f"Processing page {i}/{len(pages)}...")
    proc_img = preprocess_image(page)

    img_path = work_dir / f"page_{i:03d}.png"
    proc_img.save(img_path)

    # prepare base output name (without extension) for tesseract
    out_base = str(work_dir / f"page_{i:03d}")

    # 1) Run tesseract to produce a plain text file (out_base.txt)
    txt_cmd = [
        tesseract_exe,
        str(img_path),
        out_base,                 # tesseract will add .txt automatically
        "-l", LANG,
        "--oem", TESSERACT_OEM,
        "--psm", TESSERACT_PSM,
        "txt"
    ]

    # Note: some tesseract builds accept 'txt' implicitly; including it for clarity.
    try:
        subprocess.run(txt_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except subprocess.CalledProcessError as e:
        print(f"Warning: tesseract returned non-zero for text on page {i}: {e}")
    except FileNotFoundError:
        raise RuntimeError(f"Tesseract executable not found: {tesseract_exe}")

    # read produced text (if any)
    page_txt_path = Path(out_base + ".txt")
    if page_txt_path.exists():
        text = page_txt_path.read_text(encoding="utf-8", errors="ignore").strip()
    else:
        text = ""
    all_texts.append(text)

    # 2) Run tesseract to produce a searchable PDF for this page (out_base.pdf)
    pdf_cmd = [
        tesseract_exe,
        str(img_path),
        out_base,
        "-l", LANG,
        "--oem", TESSERACT_OEM,
        "--psm", TESSERACT_PSM,
        "pdf"
    ]
    try:
        subprocess.run(pdf_cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    except subprocess.CalledProcessError as e:
        print(f"Warning: tesseract returned non-zero for pdf on page {i}: {e}")

    page_pdf = Path(out_base + ".pdf")
    if page_pdf.exists():
        page_pdf_paths.append(page_pdf)
    else:
        print(f"Warning: expected PDF not found for page {i}: {page_pdf}")

    # Add to DOCX: page header, image, then text
    doc.add_paragraph(f"--- Page {i} ---")
    # Insert scaled image; adjust width if needed
    try:
        doc.add_picture(str(img_path), width=Inches(6))
    except Exception as e:
        print(f"Warning: couldn't insert image into docx for page {i}: {e}")
    doc.add_paragraph("")  # spacer
    doc.add_paragraph(text)

# Merge per-page PDFs into a single searchable PDF
out_pdf_path = Path(f"{OUT_PREFIX}_searchable.pdf")
try:
    try:
        from pypdf import PdfMerger
        merger = PdfMerger()
    except Exception:
        # fallback to PyPDF2's merger name differences
        from PyPDF2 import PdfFileMerger as PdfMerger
        merger = PdfMerger()
    for p in page_pdf_paths:
        merger.append(str(p))
    merger.write(str(out_pdf_path))
    merger.close()
except Exception as e:
    print(f"Warning: failed to merge PDFs: {e}")
    # As fallback, if only one page_pdf exists, copy it
    if len(page_pdf_paths) == 1:
        out_pdf_path = Path(f"{OUT_PREFIX}_searchable.pdf")
        page_pdf_paths[0].replace(out_pdf_path)
        print(f"Single-page PDF moved to {out_pdf_path}")

# Save DOCX
out_docx_path = Path(f"{OUT_PREFIX}.docx")
doc.save(str(out_docx_path))

# Save plain text as well (all pages concatenated)
out_txt_path = Path(f"{OUT_PREFIX}.txt")
with out_txt_path.open("w", encoding="utf-8") as f:
    f.write("\n\n==== PAGE BREAK ==== \n\n".join(all_texts))


# -------------------------
# Rename ocr_temp dir to <PDF_STEM>_pages and zip it
# -------------------------
pdf_stem = Path(PDF_INPUT).stem
new_dir_name = f"{pdf_stem}_pages"
new_dir_path = Path(new_dir_name)

# if a dir with that name exists, remove or rename (here: remove)
if new_dir_path.exists():
    # be careful: this will delete existing folder with same name
    print(f"Note: {new_dir_path} already exists, removing it.")
    shutil.rmtree(new_dir_path)

# move/rename work_dir -> new_dir_path
try:
    work_dir.rename(new_dir_path)
    print(f"Renamed {work_dir} -> {new_dir_path}")
except Exception as e:
    # fallback: copy tree then delete old
    print(f"Could not rename dir directly: {e}. Attempting copy...")
    shutil.copytree(work_dir, new_dir_path)
    shutil.rmtree(work_dir)
    print(f"Copied {work_dir} -> {new_dir_path} and removed old.")

# create zip archive named f"{pdf_stem}_pages.zip"
zip_base = f"{pdf_stem}_pages"
archive_path = shutil.make_archive(zip_base, 'zip', root_dir=new_dir_path)



print("Done!")
print(f"Searchable PDF: {out_pdf_path}")
print(f"DOCX (editable): {out_docx_path}")
print(f"Plain text: {out_txt_path}")
print(f"Temp files in: {work_dir}")
print(f"Created zip: {archive_path}")
