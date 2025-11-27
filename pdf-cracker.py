"""
OCR + Searchable PDF + DOCX generator (uses tesseract CLI via subprocess)

- Uses tesseract CLI to produce page_###.txt and page_###.pdf
- Merges page PDFs into a single searchable PDF (if pages produced)
- If page-PDFs are missing, will create a fallback (image-only) PDF
- Renames ocr_temp -> <input_pdf_stem>_pages and creates <input_pdf_stem>_pages.zip

Requirements:
 - Tesseract OCR installed and available (set TESSERACT_CMD if not in PATH)
 - poppler (for pdf2image)
 - pip install pdf2image python-docx pillow opencv-python pypdf
"""

import os
import subprocess
from pathlib import Path
from pdf2image import convert_from_path
from PIL import Image
import cv2
import numpy as np
from docx import Document
from docx.shared import Inches
import shutil
import glob
import sys

# ---------- Get PDF Path from user ----------
while True:
    path_input = input("Please Enter PDF Path: ").strip()
    if not path_input:
        print("[x] Empty input. Try again.")
        continue
    PDF_INPUT = Path(path_input)
    if PDF_INPUT.exists() and PDF_INPUT.is_file():
        break
    else:
        print(f"[x] PDF Not Found: {PDF_INPUT}. Enter Correct Path Please.")

# ---------- CONFIG ----------
OUT_PREFIX = "output"
DPI = 300
LANG = "eng"      # change to "fas" or "eng+fas" if you have Persian trained data
USE_DESKEW = True
POPPLER_PATH = r"H:\Repo\LittleApps\Materials\poppler\Library\bin"  # set to None if not needed
TESSERACT_CMD = None  # e.g. r"C:\Program Files\Tesseract-OCR\tesseract.exe" or None if in PATH
TESSERACT_OEM = "3"
TESSERACT_PSM = "3"
# ----------------------------

tesseract_exe = TESSERACT_CMD or "tesseract"

def check_tesseract():
    try:
        res = subprocess.run([tesseract_exe, "--version"], check=True, capture_output=True, text=True)
        print("Tesseract version:", res.stdout.splitlines()[0])
    except Exception as e:
        raise RuntimeError(f"Cannot run tesseract executable '{tesseract_exe}'. Ensure tesseract is installed and path is correct. Error: {e}")

check_tesseract()

work_dir = Path("ocr_temp")
# if exists and not empty, create unique folder to avoid overwriting
if work_dir.exists():
    if any(work_dir.iterdir()):
        # create timestamped folder to avoid accidental overwrite
        from datetime import datetime
        suffix = datetime.now().strftime("%Y%m%d_%H%M%S")
        work_dir = Path(f"ocr_temp_{suffix}")
work_dir.mkdir(parents=True, exist_ok=True)

# helper: simple deskew and denoise using OpenCV
def preprocess_image(pil_image):
    img = cv2.cvtColor(np.array(pil_image), cv2.COLOR_RGB2BGR)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    denoised = cv2.fastNlMeansDenoising(gray, None, h=10, templateWindowSize=7, searchWindowSize=21)
    th = cv2.adaptiveThreshold(denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                               cv2.THRESH_BINARY, 41, 11)
    if USE_DESKEW:
        coords = cv2.findNonZero(255 - th)
        if coords is not None:
            angle = cv2.minAreaRect(coords)[-1]
            if angle < -45:
                angle = -(90 + angle)
            else:
                angle = -angle
            (h, w) = th.shape
            center = (w // 2, h // 2)
            M = cv2.getRotationMatrix2D(center, angle, 1.0)
            rotated = cv2.warpAffine(th, M, (w, h), flags=cv2.INTER_CUBIC, borderMode=cv2.BORDER_REPLICATE)
            th = rotated
    return Image.fromarray(th)

print("Converting PDF to images...")
if POPPLER_PATH:
    pages = convert_from_path(PDF_INPUT, dpi=DPI, poppler_path=POPPLER_PATH)
else:
    pages = convert_from_path(PDF_INPUT, dpi=DPI)

doc = Document()
all_texts = []
page_pdf_paths = []

def run_tesseract_cmd(cmd):
    """Run subprocess, return (returncode, stdout, stderr)."""
    try:
        completed = subprocess.run(cmd, check=False, capture_output=True, text=True)
        return completed.returncode, completed.stdout, completed.stderr
    except FileNotFoundError as e:
        return 127, "", str(e)

for i, page in enumerate(pages, start=1):
    print(f"\n--- Processing page {i}/{len(pages)} ---")
    proc_img = preprocess_image(page)

    img_path = work_dir / f"page_{i:03d}.png"
    proc_img.save(img_path)
    print(f"Saved preprocessed image: {img_path}")

    # Use absolute output base to avoid cwd surprises
    out_base_path = (work_dir / f"page_{i:03d}").resolve()
    out_base = str(out_base_path)

    # 1) produce txt
    txt_cmd = [
        tesseract_exe,
        str(img_path),
        out_base,
        "-l", LANG,
        "--oem", TESSERACT_OEM,
        "--psm", TESSERACT_PSM,
        "txt"
    ]
    rc, so, se = run_tesseract_cmd(txt_cmd)
    if rc != 0:
        print(f"[tesseract txt] RETURN CODE {rc}\nSTDERR:\n{se.strip()}")
    else:
        print(f"[tesseract txt] OK for page {i}")

    page_txt_path = Path(out_base + ".txt")
    if page_txt_path.exists():
        try:
            text = page_txt_path.read_text(encoding="utf-8", errors="ignore").strip()
        except Exception:
            text = ""
    else:
        text = ""
        # debug: show small samples of directories
        print(f"Warning: text file not found for page {i} at {page_txt_path}")
        # show stderr for debugging
        if se:
            print("tesseract txt stderr (tail):", se[-500:])

    all_texts.append(text)

    # 2) produce pdf (searchable)
    pdf_cmd = [
        tesseract_exe,
        str(img_path),
        out_base,
        "-l", LANG,
        "--oem", TESSERACT_OEM,
        "--psm", TESSERACT_PSM,
        "pdf"
    ]
    rc, so, se = run_tesseract_cmd(pdf_cmd)
    if rc != 0:
        print(f"[tesseract pdf] RETURN CODE {rc}\nSTDERR:\n{se.strip()}")
    else:
        print(f"[tesseract pdf] OK for page {i}")

    page_pdf = Path(out_base + ".pdf")
    if page_pdf.exists():
        page_pdf_paths.append(page_pdf)
        print(f"Found page PDF: {page_pdf}")
    else:
        # try to find in cwd or work_dir (some tesseract builds write to cwd)
        cand_cwd = Path.cwd() / page_pdf.name
        cand_work = work_dir / page_pdf.name
        found = None
        if cand_cwd.exists():
            found = cand_cwd
        elif cand_work.exists():
            found = cand_work
        else:
            # broader search for any page_###.pdf recently created
            glob_candidates = list(glob.glob(f"page_{i:03d}.pdf")) + list(glob.glob(str(work_dir / f"page_{i:03d}.pdf")))
            if glob_candidates:
                found = Path(glob_candidates[0])

        if found:
            # move to expected place (work_dir)
            try:
                target = work_dir / found.name
                if found.resolve() != target.resolve():
                    shutil.move(str(found), str(target))
                page_pdf_paths.append(target)
                print(f"Moved found PDF {found} -> {target}")
            except Exception as e:
                print(f"Warning: found PDF but failed to move {found}: {e}")
        else:
            print(f"Warning: expected PDF not found for page {i}: {page_pdf}")
            if se:
                print("tesseract pdf stderr (tail):", se[-1000:])

    # Add to DOCX: header, image, text
    doc.add_paragraph(f"--- Page {i} ---")
    try:
        doc.add_picture(str(img_path), width=Inches(6))
    except Exception as e:
        print(f"Warning: couldn't insert image into docx for page {i}: {e}")
    doc.add_paragraph("")  # spacer
    doc.add_paragraph(text)

# Merge per-page PDFs into a single searchable PDF
out_pdf_path = Path(f"{OUT_PREFIX}_searchable.pdf")
if page_pdf_paths:
    print("\nMerging per-page PDFs into searchable PDF...")
    try:
        # prefer pypdf
        try:
            from pypdf import PdfMerger
            merger = PdfMerger()
        except Exception:
            from PyPDF2 import PdfFileMerger as PdfMerger
            merger = PdfMerger()

        for p in page_pdf_paths:
            merger.append(str(p))
        merger.write(str(out_pdf_path))
        merger.close()
        print(f"Searchable PDF written: {out_pdf_path}")
    except Exception as e:
        print(f"Warning: failed to merge PDFs: {e}")
        # fallback if only one pdf exists
        if len(page_pdf_paths) == 1:
            single = page_pdf_paths[0]
            try:
                shutil.copy(single, out_pdf_path)
                print(f"Single page PDF copied to {out_pdf_path}")
            except Exception as e2:
                print(f"Failed copying single page PDF: {e2}")
else:
    # No page PDFs produced — create a fallback image-only PDF so user still has output
    print("\nNo page-level PDFs were produced by tesseract. Creating fallback image-only PDF (NOT searchable).")
    try:
        img_list = [str(work_dir / f"page_{i:03d}.png") for i in range(1, len(pages)+1)]
        pil_images = [Image.open(p) .convert("RGB") for p in img_list]
        pil_images[0].save(str(out_pdf_path), save_all=True, append_images=pil_images[1:])
        print(f"Fallback (image-only) PDF written: {out_pdf_path}")
    except Exception as e:
        print(f"Failed to create fallback PDF from images: {e}")

# Save DOCX
out_docx_path = Path(f"{OUT_PREFIX}.docx")
try:
    doc.save(str(out_docx_path))
    print(f"DOCX saved: {out_docx_path}")
except Exception as e:
    print(f"Failed saving DOCX: {e}")

# Save concatenated plain text
out_txt_path = Path(f"{OUT_PREFIX}.txt")
try:
    with out_txt_path.open("w", encoding="utf-8") as f:
        f.write("\n\n==== PAGE BREAK ==== \n\n".join(all_texts))
    print(f"Plain text saved: {out_txt_path}")
except Exception as e:
    print(f"Failed saving plain text: {e}")

# -------------------------
# Rename ocr_temp dir to <PDF_STEM>_pages and zip it
# -------------------------
pdf_stem = Path(PDF_INPUT).stem
new_dir_name = f"{pdf_stem}_pages"
new_dir_path = Path(new_dir_name)

if new_dir_path.exists():
    print(f"Note: {new_dir_path} already exists, removing it.")
    shutil.rmtree(new_dir_path)

try:
    # work_dir may have been changed name earlier; ensure we rename the correct folder
    current_work_dir = work_dir
    if current_work_dir.exists():
        current_work_dir.rename(new_dir_path)
        print(f"Renamed {current_work_dir} -> {new_dir_path}")
    else:
        print(f"Warning: expected work dir {current_work_dir} not found to rename.")
        # try to create copy from any ocr_temp* created
        possible = sorted(glob.glob("ocr_temp*"), reverse=True)
        if possible:
            shutil.copytree(possible[0], new_dir_path)
            print(f"Copied {possible[0]} -> {new_dir_path}")
except Exception as e:
    print(f"Could not rename dir directly: {e}. Attempting copy fallback...")
    try:
        shutil.copytree(work_dir, new_dir_path)
        shutil.rmtree(work_dir)
        print(f"Copied {work_dir} -> {new_dir_path} and removed old.")
    except Exception as e2:
        print(f"Fallback copy also failed: {e2}")

# create zip archive named f"{pdf_stem}_pages.zip"
zip_base = f"{pdf_stem}_pages"
try:
    archive_path = shutil.make_archive(zip_base, 'zip', root_dir=new_dir_path)
    print(f"Created zip: {archive_path}")
except Exception as e:
    print(f"Failed creating zip archive: {e}")

print("\nAll done.")
print(f"Searchable PDF (or fallback): {out_pdf_path.resolve()}")
print(f"DOCX (editable): {out_docx_path.resolve() if out_docx_path.exists() else 'NOT_CREATED'}")
print(f"Plain text: {out_txt_path.resolve() if out_txt_path.exists() else 'NOT_CREATED'}")
print(f"Pages folder (renamed): {new_dir_path.resolve() if new_dir_path.exists() else work_dir.resolve()}")
if 'archive_path' in locals():
    print(f"ZIP: {Path(archive_path).resolve()}")
